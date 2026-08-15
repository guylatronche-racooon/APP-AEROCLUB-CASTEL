#!/usr/bin/env python3
"""Generate the editable user manual and its illustrated assets.

Design basis: compact_reference_guide, with one named override for a French
software operator manual (A4, Arial, ACJD colours, restrained title page).
"""

from __future__ import annotations

import argparse
import json
import math
import os
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "manuel-utilisation-outils-de-vol.docx"
ASSET_DIR = ROOT / "tmp" / "manual-assets"

NAVY = "073B4C"
TEAL = "087F70"
CYAN = "129DD0"
INK = "172A34"
MUTED = "627783"
LINE = "CFDCE1"
PALE_BLUE = "EAF6FA"
PALE_YELLOW = "FFF5CF"
PALE_RED = "FBEAE7"
PALE_GREEN = "EAF7F2"
WHITE = "FFFFFF"

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_LEFT_CM = 1.65
MARGIN_RIGHT_CM = 1.65
MARGIN_TOP_CM = 1.65
MARGIN_BOTTOM_CM = 1.55
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM
CONTENT_WIDTH_DXA = round(CONTENT_WIDTH_CM / 2.54 * 1440)
TABLE_INDENT_DXA = 120

SCREENSHOTS = {
    "hub": (ROOT / "upload" / "F97F92BB-9D63-48A7-8683-3ABF33CB6B66.png", (0, 125, 828, 1720)),
    "balance": (ROOT / "upload" / "6AC0DA4A-759A-45F9-9937-7B1CAFC5FA6E.png", (0, 225, 828, 1580)),
    "balance_result": (ROOT / "upload" / "D4277F8D-630C-4CFE-91AA-2F8634786C4F.png", (0, 70, 828, 1260)),
    "density": (ROOT / "upload" / "8B5E86CD-C080-4702-B1AA-21844286E9D5.png", (0, 135, 828, 1720)),
    "performance": (ROOT / "upload" / "6127A0AF-64C2-49CB-A60B-152525FFEDD2.png", (0, 310, 828, 1730)),
    "performance_result": (ROOT / "upload" / "194A71F2-FA20-4AAF-94F0-8E6D04DC58AE.png", (0, 95, 828, 1600)),
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_font(run, name="Arial", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge in edges.items():
        tag = "start" if edge_name == "left" else "end" if edge_name == "right" else edge_name
        element = borders.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            borders.append(element)
        for key, value in edge.items():
            element.set(qn(f"w:{key}"), str(value))


def configure_table(table, widths_cm, header=True, quiet=False):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    widths = [round(value / CONTENT_WIDTH_CM * CONTENT_WIDTH_DXA) for value in widths_cm]
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        if row_index == 0 and header:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border = {"val": "single", "sz": "4", "color": LINE}
            if quiet:
                border = {"val": "single", "sz": "3", "color": "DFE7EA"}
            set_cell_border(cell, top=border, bottom=border, left=border, right=border)
            if row_index == 0 and header:
                shade_cell(cell, PALE_BLUE)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, size=8.5, color=NAVY, bold=True)


def table_set_text(cell, text, bold=False, color=INK, size=8.7):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    set_font(r, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text, url, color=TEAL):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text_node, end):
        run._r.append(element)


def prepare_screenshots():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    prepared = {}
    for key, (source, crop_box) in SCREENSHOTS.items():
        if not source.exists():
            continue
        target = ASSET_DIR / f"{key}.jpg"
        with Image.open(source) as image:
            image = image.convert("RGB").crop(crop_box)
            if image.width > 620:
                height = round(image.height * 620 / image.width)
                image = image.resize((620, height), Image.Resampling.LANCZOS)
            image.save(target, format="JPEG", quality=68, optimize=True, progressive=True)
        prepared[key] = target

    icon_source = ROOT / "app" / "icons" / "icon-512.png"
    if icon_source.exists():
        icon_target = ASSET_DIR / "app-icon.jpg"
        with Image.open(icon_source) as image:
            image = image.convert("RGB").resize((240, 240), Image.Resampling.LANCZOS)
            image.save(icon_target, format="JPEG", quality=82, optimize=True, progressive=True)
        prepared["icon"] = icon_target
    return prepared


def paragraph(doc, text="", *, style=None, bold_prefix=None, italic=False, color=INK, size=None, align=None, after=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=size, color=color, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="Manual Bullet" if level == 0 else "Manual Bullet 2")
    p.add_run(text)
    return p


def add_step(doc, title, text):
    p = doc.add_paragraph(style="Manual Step")
    previous = doc.paragraphs[-2] if len(doc.paragraphs) >= 2 else None
    if previous is None or previous.style.name != "Manual Step":
        numbering = doc.part.numbering_part.element
        next_num_id = max([int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))] + [0]) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(doc._manual_step_abstract_id))
        num.append(abstract_ref)
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
        numbering.append(num)
        doc._manual_step_num_id = next_num_id
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(doc._manual_step_num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    r = p.add_run(title + ". ")
    set_font(r, size=9.7, color=INK, bold=True)
    r = p.add_run(text)
    set_font(r, size=9.7, color=INK)
    return p


def add_callout(doc, label, text, kind="note"):
    fill = {"note": PALE_BLUE, "attention": PALE_YELLOW, "reserve": PALE_RED, "ok": PALE_GREEN}[kind]
    accent = {"note": CYAN, "attention": "C79D16", "reserve": "AA4237", "ok": TEAL}[kind]
    table = doc.add_table(rows=1, cols=1)
    configure_table(table, [CONTENT_WIDTH_CM], header=False, quiet=True)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": accent}, top={"val": "single", "sz": "2", "color": fill}, bottom={"val": "single", "sz": "2", "color": fill}, right={"val": "single", "sz": "2", "color": fill})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper())
    set_font(r, size=8.5, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r, size=9.1, color=INK)
    paragraph(doc, "", after=2)


def add_figure(doc, path, caption, width_cm=7.2, page_break_before=False):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break_before
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_together = True
    r = cap.add_run(caption)
    set_font(r, size=7.8, color=MUTED, italic=True)


def add_two_column_figure(doc, figure_path, caption, right_title, right_paragraphs, width_left=7.0):
    table = doc.add_table(rows=1, cols=2)
    configure_table(table, [7.5, CONTENT_WIDTH_CM - 7.5], header=False, quiet=True)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
        set_cell_margins(cell, top=20, start=40, bottom=20, end=100)
    left = table.cell(0, 0)
    left.text = ""
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(figure_path), width=Cm(width_left))
    p = left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_font(r, size=7.5, color=MUTED, italic=True)
    right = table.cell(0, 1)
    right.text = ""
    p = right.paragraphs[0]
    r = p.add_run(right_title)
    set_font(r, size=11, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(5)
    for item in right_paragraphs:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_font(r, size=9.2, color=INK)


def heading(doc, text, level=1, page_break=False):
    p = doc.add_heading(text, level=level)
    if page_break:
        p.paragraph_format.page_break_before = True
    return p


def add_data_table(doc, headers, rows, widths_cm, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table_set_text(table.rows[0].cells[index], value, bold=True, color=NAVY, size=8.3)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            table_set_text(cells[index], value, size=font_size)
    configure_table(table, widths_cm, header=True)
    paragraph(doc, "", after=2)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 8),
        ("Heading 1", 18, NAVY, 16, 9),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 10.8, NAVY, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Title" or True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Manual Bullet" not in styles:
        bullet = styles.add_style("Manual Bullet", 1)
    else:
        bullet = styles["Manual Bullet"]
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Cm(0.75)
    bullet.paragraph_format.first_line_indent = Cm(-0.38)
    bullet.paragraph_format.space_after = Pt(3)
    bullet.paragraph_format.line_spacing = 1.14
    bullet.font.name = "Arial"
    bullet.font.size = Pt(9.5)

    if "Manual Bullet 2" not in styles:
        bullet2 = styles.add_style("Manual Bullet 2", 1)
    else:
        bullet2 = styles["Manual Bullet 2"]
    bullet2.base_style = normal
    bullet2.paragraph_format.left_indent = Cm(1.25)
    bullet2.paragraph_format.first_line_indent = Cm(-0.35)
    bullet2.paragraph_format.space_after = Pt(2)
    bullet2.font.name = "Arial"
    bullet2.font.size = Pt(9.3)

    if "Manual Step" not in styles:
        step = styles.add_style("Manual Step", 1)
    else:
        step = styles["Manual Step"]
    step.base_style = normal
    step.paragraph_format.left_indent = Cm(0.85)
    step.paragraph_format.first_line_indent = Cm(-0.55)
    step.paragraph_format.space_after = Pt(4)
    step.paragraph_format.line_spacing = 1.14
    step.font.name = "Arial"
    step.font.size = Pt(9.7)

    # Real numbering definitions: bullet marker and decimal step number.
    numbering = doc.part.numbering_part.element
    abstract_start = max([int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))] + [0]) + 1
    num_start = max([int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))] + [0]) + 1

    def add_num(abstract_id, num_id, fmt, text_value, left, hanging, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    bullet_id = add_num(abstract_start, num_start, "bullet", "•", 540, 270, "Arial")
    step_id = add_num(abstract_start + 1, num_start + 1, "decimal", "%1.", 720, 360, "Arial")
    bullet2_id = add_num(abstract_start + 2, num_start + 2, "bullet", "–", 900, 270, "Arial")

    def attach_num(style, num_id):
        p_pr = style._element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num)

    attach_num(bullet, bullet_id)
    attach_num(bullet2, bullet2_id)
    attach_num(step, step_id)
    doc._manual_step_abstract_id = abstract_start + 1
    doc._manual_step_num_id = step_id


def configure_section(section):
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    table = header.add_table(rows=1, cols=2, width=Cm(CONTENT_WIDTH_CM))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_width(table.cell(0, 0), round(CONTENT_WIDTH_DXA * 0.55))
    set_cell_width(table.cell(0, 1), round(CONTENT_WIDTH_DXA * 0.45))
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=35, start=0, end=0)
        set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": LINE})
    left = table.cell(0, 0).paragraphs[0]
    r = left.add_run("OUTILS DE VOL")
    set_font(r, size=8.2, color=NAVY, bold=True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right.add_run("Manuel utilisateur - version 0.9")
    set_font(r, size=8.2, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Pt(0)
    table = footer.add_table(rows=1, cols=2, width=Cm(CONTENT_WIDTH_CM))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_cell_width(table.cell(0, 0), round(CONTENT_WIDTH_DXA * 0.7))
    set_cell_width(table.cell(0, 1), round(CONTENT_WIDTH_DXA * 0.3))
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=35, bottom=0, start=0, end=0)
        set_cell_border(cell, top={"val": "single", "sz": "5", "color": LINE})
    left = table.cell(0, 0).paragraphs[0]
    r = left.add_run("Document de travail - validation club requise")
    set_font(r, size=7.8, color=MUTED)
    add_page_number(table.cell(0, 1).paragraphs[0])


def page_break(doc):
    doc.add_page_break()


def section_intro(doc, number, title, lead, new_page=True):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = new_page
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(number)
    set_font(r, size=9, color=CYAN, bold=True)
    r.add_break()
    r = p.add_run(title)
    set_font(r, size=18, color=NAVY, bold=True)
    p = paragraph(doc, lead, size=10.2, color=MUTED, after=9)
    p.paragraph_format.line_spacing = 1.16


def create_manual(toc_pages):
    images = prepare_screenshots()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    configure_section(section)

    # Cover: editorial_cover structure, deliberately restrained for an operator manual.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(18)
    p.add_run().add_picture(str(images.get("icon", ROOT / "app" / "icons" / "icon-512.png")), width=Cm(2.8))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AÉROCLUB JEAN-DOUDIÈS")
    set_font(r, size=10, color=CYAN, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Outils de vol")
    set_font(r, size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Manuel utilisateur et méthodes de calcul")
    set_font(r, size=16, color=TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Check-lists - Masse et centrage - Altitude-densité - Performances au décollage")
    set_font(r, size=10.5, color=MUTED)

    add_callout(doc, "Statut du document", "Version de travail destinée à la relecture du club. Elle décrit le fonctionnement de l'application au 15 août 2026. Les données et méthodes signalées comme provisoires ne doivent pas être considérées comme approuvées.", "attention")

    table = doc.add_table(rows=4, cols=2)
    metadata = [
        ("Référence", "MAN-ODV-001"),
        ("Version", "0.9"),
        ("Date", "15/08/2026"),
        ("Application", "app-aeroclub-castel.vercel.app"),
    ]
    for row, values in zip(table.rows, metadata):
        table_set_text(row.cells[0], values[0], bold=True, color=NAVY, size=8.5)
        table_set_text(row.cells[1], values[1], size=8.5)
    configure_table(table, [4.0, CONTENT_WIDTH_CM - 4.0], header=False, quiet=True)
    paragraph(doc, "", after=10)
    paragraph(doc, "Ce manuel ne remplace ni le manuel de vol approuvé, ni les documents officiels de l'aéronef, ni les informations aéronautiques et météorologiques en vigueur.", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)
    heading(doc, "Suivi du document", 1)
    paragraph(doc, "Le manuel et l'application doivent évoluer ensemble. Toute modification d'une donnée avion, d'une table ou d'une règle de calcul doit entraîner une nouvelle révision de ce document.")
    add_data_table(doc, ["Version", "Date", "Objet", "Statut"], [["0.9", "15/08/2026", "Première édition détaillée", "À valider"]], [2.2, 2.8, 8.0, CONTENT_WIDTH_CM - 13.0])
    heading(doc, "Diffusion et responsabilité", 2)
    paragraph(doc, "L'application est une aide à la préparation. Le commandant de bord reste responsable des données saisies, de la consultation des documents applicables et de la décision de vol. Une information préremplie doit être contrôlée avant validation.")
    heading(doc, "Conventions utilisées", 2)
    add_data_table(doc, ["Repère", "Signification"], [
        ("Obligatoire", "L'action doit être confirmée pour qu'un calcul de performance soit présenté."),
        ("Facultatif", "L'option complète le calcul sans être nécessaire à son exécution."),
        ("CALCUL BLOQUÉ", "Une donnée ou une confirmation indispensable manque ou est incohérente."),
        ("HORS TABLE", "La configuration sort du domaine publié ou utilise une condition sans correction validée."),
        ("À valider", "La source, l'applicabilité ou la transcription n'a pas encore été approuvée par le club."),
    ], [4.2, CONTENT_WIDTH_CM - 4.2])
    add_callout(doc, "Réserve générale", "Une marge positive affichée par l'application ne constitue pas, à elle seule, une autorisation ou une décision de décollage. Obstacles, trajectoire, gradient de montée, vent traversier, rafales, état réel de piste et limitations opérationnelles restent hors du verdict.", "reserve")

    page_break(doc)
    heading(doc, "Sommaire", 1)
    toc_items = [
        ("1", "Objet, périmètre et conditions d'emploi"),
        ("2", "Prise en main"),
        ("3", "Check-lists"),
        ("4", "Masse et centrage"),
        ("5", "Calculateur d'altitude-densité"),
        ("6", "Performances au décollage"),
        ("7", "Terrain, météo, piste et vent"),
        ("8", "Lecture des résultats et des messages"),
        ("9", "Données locales, mode hors ligne et mises à jour"),
        ("10", "Incidents et reprise en saisie manuelle"),
        ("Annexe A", "Méthodes de calcul"),
        ("Annexe B", "Données intégrées par avion"),
        ("Annexe C", "Réserves d'emploi et points à valider"),
        ("Annexe D", "Glossaire"),
    ]
    table = doc.add_table(rows=0, cols=3)
    for key, title in toc_items:
        row = table.add_row()
        table_set_text(row.cells[0], key, bold=True, color=CYAN, size=9)
        table_set_text(row.cells[1], title, size=9.3)
        page_value = toc_pages.get(title, "-") if toc_pages else "-"
        table_set_text(row.cells[2], page_value, bold=True, color=NAVY, size=9)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    configure_table(table, [2.6, CONTENT_WIDTH_CM - 4.0, 1.4], header=False, quiet=True)
    paragraph(doc, "Les numéros de page correspondent à la présente version du PDF.", size=8, color=MUTED, italic=True)

    section_intro(doc, "01", "Objet, périmètre et conditions d'emploi", "Le manuel décrit les fonctions de l'application, le parcours conseillé et les règles appliquées par les calculateurs.")
    heading(doc, "1.1 Fonctions couvertes", 2)
    paragraph(doc, "L'application regroupe quatre familles d'outils : les check-lists propres aux avions, le calcul de masse et de centrage, un calculateur atmosphérique d'altitude-densité et l'évaluation des distances de décollage à partir de tables publiées ou photographiées.")
    add_bullet(doc, "Sélection d'un avion et accès à ses outils disponibles.")
    add_bullet(doc, "Saisie du chargement, calcul des moments et contrôle de l'enveloppe numérisée.")
    add_bullet(doc, "Chargement facultatif des données terrain et d'un METAR récent.")
    add_bullet(doc, "Proposition d'un QFU à partir du vent moyen, avec possibilité de le modifier.")
    add_bullet(doc, "Interpolation stricte des tables de décollage, sans extrapolation hors domaine.")
    add_bullet(doc, "Consultation des méthodes, limites et scans sources.")
    heading(doc, "1.2 Fonctions non couvertes", 2)
    paragraph(doc, "Le calcul de piste ne traite que le roulement et la distance jusqu'au passage des 15 m / 50 ft. Il ne calcule ni la trajectoire après cet écran, ni le franchissement des obstacles, ni le gradient de montée, ni une limitation de vent traversier. Il ne prépare pas le carburant réglementaire et ne calcule aucune réserve de route, de dégagement ou finale.")
    heading(doc, "1.3 Hiérarchie documentaire", 2)
    add_data_table(doc, ["Ordre", "Document ou information"], [
        ("1", "Manuel de vol approuvé, suppléments applicables et fiche de pesée en vigueur."),
        ("2", "AIP, VAC, NOTAM, briefing météo et informations opérationnelles du jour."),
        ("3", "Consignes du club et décision du commandant de bord."),
        ("4", "Préremplissages, calculs et avertissements de l'application."),
    ], [2.0, CONTENT_WIDTH_CM - 2.0])
    add_callout(doc, "Principe", "En cas d'écart entre l'application et une source officielle, la source officielle prévaut. Le calcul doit alors être interrompu et l'écart signalé.", "reserve")

    section_intro(doc, "02", "Prise en main", "Le parcours normal suit l'ordre chargement, conditions du jour, piste, puis résultat. L'utilisateur peut revenir à l'étape précédente sans perdre sa saisie.")
    if "hub" in images:
        add_two_column_figure(doc, images["hub"], "Figure 1 - Écran d'un avion (capture représentative).", "Parcours conseillé", [
            "1. Choisir l'avion utilisé.",
            "2. Exécuter la check-list lorsque la fiche est disponible.",
            "3. Renseigner Masse & centrage.",
            "4. Passer directement aux Performances.",
            "5. Contrôler le terrain, la météo et la piste avant de confirmer.",
        ], width_left=5.5)
    heading(doc, "2.1 Accéder à l'application", 2)
    add_step(doc, "Ouvrir", "Saisir app-aeroclub-castel.vercel.app dans Safari, Chrome ou un navigateur équivalent.")
    add_step(doc, "Lire l'avertissement", "À chaque ouverture, prendre connaissance du rappel documentaire puis sélectionner « J'ai compris ».")
    add_step(doc, "Choisir l'avion", "Sélectionner l'immatriculation. Les outils non documentés pour cet avion restent indisponibles.")
    heading(doc, "2.2 Installer sur l'écran d'accueil d'un iPhone", 2)
    add_step(doc, "Dans Safari", "ouvrir le menu Partager.")
    add_step(doc, "Ajouter", "choisir « Sur l'écran d'accueil », puis confirmer le nom de l'application.")
    add_step(doc, "Après une mise à jour", "fermer complètement l'application et la rouvrir. Si l'ancienne version reste affichée, ouvrir une fois le site dans Safari avec le réseau disponible.")
    heading(doc, "2.3 Navigation et sauvegarde", 2)
    paragraph(doc, "Les valeurs de chargement sont conservées localement sur l'appareil. Les confirmations de chargement, de météo et de piste ne sont pas restaurées comme valides après un rechargement. Les valeurs du jour doivent donc être revérifiées.")
    add_callout(doc, "Bon usage", "Utiliser « Nouveau vol » ou « Réinitialiser » au début d'une nouvelle préparation. Ne pas reprendre une confirmation provenant d'un vol précédent.", "attention")

    section_intro(doc, "03", "Check-lists", "Les check-lists reprennent les fiches disponibles pour chaque avion. Elles servent d'aide à l'exécution et signalent les items restant à traiter.")
    heading(doc, "3.1 Exécuter une phase", 2)
    add_step(doc, "Ouvrir", "sélectionner Check-list depuis la fiche de l'avion.")
    add_step(doc, "Choisir la phase", "ouvrir le menu des phases et sélectionner la phase correspondant à la situation de vol.")
    add_step(doc, "Lire puis agir", "effectuer l'action, contrôler la réponse et cocher l'item seulement après exécution.")
    add_step(doc, "Contrôler l'avancement", "le compteur de phase indique le nombre d'items cochés. Une phase incomplète reste visible.")
    heading(doc, "3.2 Repères avion", 2)
    paragraph(doc, "Le panneau « Repères avion » rassemble quelques valeurs reprises de la fiche photographiée. Il ne constitue pas un extrait approuvé du manuel de vol. Toute valeur doit être confirmée dans la documentation applicable.")
    heading(doc, "3.3 Nouvelle préparation", 2)
    paragraph(doc, "La progression ne doit pas être interprétée comme un enregistrement durable. Après un rechargement, une fermeture complète ou la commande « Nouveau vol », les cases repartent vides.")
    add_callout(doc, "Réserve", "Des items restent soumis à relecture lorsque le document d'origine est incomplet ou peu lisible. Un libellé marqué « à confirmer » ne doit pas être utilisé comme instruction approuvée.", "reserve")

    section_intro(doc, "04", "Masse et centrage", "Le calcul reprend la masse à vide, les bras et l'enveloppe associés à l'avion sélectionné. Le carburant est initialisé au plein pour être ensuite ramené à la quantité réellement embarquée.")
    if "balance" in images:
        add_two_column_figure(doc, images["balance"], "Figure 2 - Saisie du chargement.", "Saisie", [
            "Les occupants et les bagages sont saisis en kilogrammes.",
            "Le carburant est saisi en litres totaux présents à bord.",
            "Les quantités préremplies correspondent au plein des réservoirs documentés.",
            "Une valeur rouge, absente ou hors limite doit être corrigée.",
        ])
    heading(doc, "4.1 Procédure", 2)
    add_step(doc, "Vérifier l'avion", "contrôler l'immatriculation et la date de la fiche de pesée affichée en haut de l'écran.")
    add_step(doc, "Renseigner les personnes", "saisir chaque occupant dans le poste correspondant. Le poste pilote doit contenir une masse strictement positive.")
    add_step(doc, "Renseigner les bagages", "saisir la masse réellement embarquée et vérifier les limites dans le manuel de vol.")
    add_step(doc, "Ajuster le carburant", "partir des pleins préremplis et remplacer chaque quantité par les litres réellement présents au départ.")
    add_step(doc, "Lire le résultat", "contrôler la masse totale, le centre de gravité, la marge de masse et la position du point dans l'enveloppe.")
    add_step(doc, "Poursuivre", "si le chargement est conforme, sélectionner « Chargement vérifié - passer aux performances ». Cette action vaut confirmation explicite du chargement.")
    heading(doc, "4.2 Particularités carburant", 2)
    add_data_table(doc, ["Avion", "Valeur initiale", "Traitement"], [
        ("F-HDLT / F-HDLV", "120 l", "Masse ajoutée = litres x 0,72 kg/l."),
        ("F-GGHL - principal", "110 l", "Les 10 l inutilisables sont déjà dans la masse à vide. Masse ajoutée = max(0, litres - 10) x 0,72."),
        ("F-GGHL - ailes", "80 l", "Masse ajoutée = litres x 0,72 kg/l."),
    ], [4.0, 3.2, CONTENT_WIDTH_CM - 7.2], font_size=8.2)
    paragraph(doc, "Le bouton « Réinitialiser » remet les réservoirs au plein et efface les autres postes. Une quantité déjà ajustée reste conservée lors d'une navigation normale entre les écrans.", size=9.2)
    heading(doc, "4.3 Lecture du tableau et du graphique", 2)
    paragraph(doc, "Chaque ligne affiche la masse retenue, le bras et le moment. Le total utilise le moment à vide exact publié. Le centre de gravité est obtenu par division du moment total par la masse totale. Le graphique place ensuite le point dans l'enveloppe numérisée.")
    if "balance_result" in images:
        add_figure(doc, images["balance_result"], "Figure 3 - Exemple de résultat dans l'enveloppe numérisée. La capture ne vaut pas validation du chargement.", width_cm=9.1)
    add_callout(doc, "Contrôle obligatoire", "« Dans l'enveloppe numérisée » signifie que le point se trouve dans le contour transcrit dans l'application. Les limites exactes et la catégorie applicable doivent rester contrôlées dans le manuel de vol et la fiche de pesée en vigueur.", "reserve")

    section_intro(doc, "05", "Calculateur d'altitude-densité", "Ce module est un calcul atmosphérique général. Il n'est pas rattaché à un avion et peut être utilisé par saisie manuelle ou à partir d'un terrain et d'un METAR.")
    if "density" in images:
        add_two_column_figure(doc, images["density"], "Figure 4 - Données d'entrée du calculateur.", "Deux modes", [
            "Mode manuel : saisir altitude terrain, QNH et température.",
            "Mode METAR : cocher l'option, saisir l'OACI puis charger les données.",
            "Terrain et météo sont traités séparément : l'un peut fonctionner sans l'autre.",
            "Le point de rosée n'est nécessaire que pour l'option humidité.",
        ], width_left=4.8)
    heading(doc, "5.1 Saisie manuelle", 2)
    add_step(doc, "Altitude terrain", "saisir l'élévation publiée de l'aérodrome en pieds.")
    add_step(doc, "QNH", "saisir la pression réduite au niveau de la mer en hPa.")
    add_step(doc, "Température", "saisir la température extérieure observée au terrain en degrés Celsius.")
    add_step(doc, "Point de rosée", "le renseigner si l'effet de l'humidité doit être estimé.")
    heading(doc, "5.2 Utiliser les données du METAR", 2)
    add_step(doc, "Activer", "cocher « Utiliser les données du METAR ».")
    add_step(doc, "Saisir l'OACI", "entrer les quatre lettres de l'aérodrome. Le champ blanc constitue la zone modifiable.")
    add_step(doc, "Charger", "sélectionner « Utiliser le METAR ». L'altitude vient du jeu terrain local ; le QNH, la température et le point de rosée viennent d'une observation récente lorsque celle-ci est exploitable.")
    add_step(doc, "Contrôler", "lire le code de la station, l'heure de l'observation et l'éventuelle mention « station voisine » avant d'utiliser les valeurs.")
    heading(doc, "5.3 Option humidité", 2)
    paragraph(doc, "L'option « Tenir compte de l'humidité » calcule la densité réelle d'un mélange d'air sec et de vapeur d'eau à partir de la pression, de la température et du point de rosée. Elle est indépendante de tout avion et n'est donc rattachée à aucune table constructeur.")
    paragraph(doc, "L'écran affiche, selon le cas, l'altitude-pression, la température ISA, l'écart ISA, l'humidité estimée, l'altitude-densité sèche et l'effet de l'humidité.")
    add_callout(doc, "Réserve", "Le calcul atmosphérique de l'air humide est physiquement général. Son emploi pour modifier une distance de décollage n'est traité que dans l'écran Performances, comme comparaison théorique, et doit être vérifié dans le manuel de l'avion.", "attention")

    section_intro(doc, "06", "Performances au décollage", "L'écran combine le chargement de l'avion, les conditions atmosphériques, la piste choisie et la table de performance disponible.", new_page=False)
    heading(doc, "6.1 Conditions préalables", 2)
    paragraph(doc, "Le calcul n'est présenté que si les trois contrôles obligatoires sont en cours de validité : chargement, conditions du jour, piste et distances. Chaque confirmation expire après 60 minutes. Une confirmation météo issue d'un METAR expire également lorsque l'observation dépasse deux heures.")
    if "performance" in images:
        add_two_column_figure(doc, images["performance"], "Figure 5 - Terrain, METAR et sélection de piste.", "Ordre de travail", [
            "1. Confirmer le chargement.",
            "2. Saisir le terrain et charger Terrain + METAR.",
            "3. Vérifier la piste proposée.",
            "4. Contrôler QNH, température, vent, pente, surface, TORA et TODA.",
            "5. Cocher les confirmations obligatoires.",
        ], width_left=6.1)
    heading(doc, "6.2 Procédure détaillée", 2)
    add_step(doc, "Chargement", "venir de l'écran Masse et centrage ou cocher la confirmation après avoir revérifié chaque poste.")
    add_step(doc, "Terrain", "saisir l'indicatif OACI dans le champ blanc puis sélectionner « Terrain + METAR ».")
    add_step(doc, "Piste", "contrôler le QFU proposé et le changer si les informations opérationnelles imposent un autre sens.")
    add_step(doc, "Météo", "contrôler l'heure et la station du METAR. Confirmer les conditions seulement après comparaison avec les observations disponibles au terrain.")
    add_step(doc, "Distances", "vérifier TORA et TODA sur la VAC, l'AIP et les NOTAM. Corriger les valeurs si nécessaire.")
    add_step(doc, "Surface et pente", "choisir l'état réel de la piste et vérifier la pente. Une valeur 0 % marquée « non publiée » est une valeur par défaut, pas une donnée officielle.")
    add_step(doc, "Vent", "contrôler la composante préremplie. Passer en saisie manuelle si la valeur retenue ne représente pas les conditions au décollage.")
    add_step(doc, "Marge", "renseigner la marge additionnelle décidée selon les consignes du club et l'appréciation du pilote. La valeur initiale 0 % n'est pas une recommandation.")
    add_step(doc, "Confirmer", "cocher « Piste et distances vérifiées », puis lire le résultat et toutes les limites actives.")
    heading(doc, "6.3 Surface de piste", 2)
    add_data_table(doc, ["Choix", "Traitement actuel"], [
        ("Dur, sec", "Colonne piste dure de la table."),
        ("Herbe, sèche", "Colonne herbe de la table."),
        ("Dur, mouillé", "HORS TABLE tant qu'aucun facteur applicable n'est validé."),
        ("Herbe, mouillée", "HORS TABLE tant qu'aucun facteur applicable n'est validé."),
    ], [4.2, CONTENT_WIDTH_CM - 4.2])
    paragraph(doc, "Les options sont filtrées selon la surface associée au QFU sélectionné. Une piste dure et une bande en herbe sont traitées comme des entrées distinctes lorsque les distances diffèrent.")
    heading(doc, "6.4 Comparaison avec l'air humide", 2)
    paragraph(doc, "Cette option est facultative. L'application convertit la densité humide en altitude-pression sèche équivalente, recalcule la table et conserve la plus grande des distances sèche et humide. Si cette conversion sort du domaine de la table, le résultat complet passe hors table.")
    add_callout(doc, "Portée", "L'option humidité ne peut jamais améliorer un résultat. Elle ne rend pas acceptable une configuration déjà défavorable dans la méthode publiée.", "attention")

    section_intro(doc, "07", "Terrain, météo, piste et vent", "Les préremplissages réduisent la saisie, mais ne déterminent pas la piste en service et ne remplacent pas les informations du jour.")
    heading(doc, "7.1 Données terrain", 2)
    paragraph(doc, "Le jeu local contient les aérodromes métropolitains du cycle AIRAC 08/26. Il reprend notamment l'altitude, les QFU, les surfaces, les dimensions, les relèvements vrais et les altitudes de seuil. Les distances déclarées TORA et TODA ne sont remplies que lorsqu'elles ont été intégrées depuis une source identifiée.")
    paragraph(doc, "Une absence de donnée laisse le champ modifiable. La longueur physique d'une piste ne doit pas être confondue avec une TORA ou une TODA publiée.")
    heading(doc, "7.2 METAR et station voisine", 2)
    paragraph(doc, "Le relais interroge Aviation Weather Center. Un METAR n'est utilisé que si la station retournée est celle demandée, si l'heure est cohérente, si l'observation a moins de deux heures et si QNH et température sont présents. Une station voisine prévue dans le jeu terrain est indiquée explicitement.")
    paragraph(doc, "Le TAF est affiché pour information. Il n'alimente jamais le calcul instantané.")
    heading(doc, "7.3 Proposition de piste", 2)
    paragraph(doc, "Lorsque le METAR contient une direction et une vitesse de vent moyen, l'application calcule la composante longitudinale sur chaque QFU documenté et propose celui qui donne la plus forte composante de face. En cas d'égalité pratique, aucune proposition automatique n'est imposée.")
    paragraph(doc, "Exemple : à Quimper (LFRQ), un vent moyen du 350° conduit à proposer le QFU 27 plutôt que le QFU 09. Le choix reste modifiable et doit être confronté à la piste en service, à la VAC et aux NOTAM.")
    heading(doc, "7.4 Composante de vent", 2)
    paragraph(doc, "La composante affichée utilise le vent moyen et le relèvement vrai du QFU : composante = vitesse x cos(différence angulaire). Une valeur positive est un vent de face ; une valeur négative est un vent arrière.")
    add_bullet(doc, "Le secteur variable est affiché sous forme d'une plage de composantes.")
    add_bullet(doc, "La rafale est affichée séparément et n'est pas intégrée silencieusement au champ.")
    add_bullet(doc, "Si le secteur variable traverse vent de face et vent arrière, 0 kt est proposé par défaut ; le pilote doit confirmer ou ajuster.")
    add_bullet(doc, "Si le vent est VRB ou incomplet, la composante reste à 0 kt par défaut et doit être saisie manuellement.")
    add_callout(doc, "Important", "La meilleure composante longitudinale ne désigne pas la piste en service. Elle ne tient compte ni du vent traversier maximal, ni des restrictions locales, ni des circuits, ni d'une fermeture ou d'une distance réduite.", "reserve")
    heading(doc, "7.5 Pente", 2)
    paragraph(doc, "La pente n'est pas un champ explicitement publié dans l'export SIA utilisé. Lorsqu'elle manque, l'application propose 0 % et indique « non publiée ». Cette valeur doit être vérifiée sur la documentation du terrain. Une pente montante rend actuellement le calcul hors table en l'absence de correction validée ; aucun gain de pente descendante n'est inventé.")
    heading(doc, "7.6 Terrains de montagne", 2)
    paragraph(doc, "Pour un altiport, une altisurface ou un terrain de montagne, consulter la documentation propre au terrain, les procédures, les limitations et la réglementation applicables aux conditions particulières. Le calcul générique de longueur de piste ne couvre pas ces exigences.")

    section_intro(doc, "08", "Lecture des résultats et des messages", "L'application distingue un calcul incomplet, une condition hors table et un résultat chiffré. Le détail sous le bandeau explique toujours la cause retenue.")
    heading(doc, "8.1 États possibles", 2)
    add_data_table(doc, ["État", "Sens", "Action"], [
        ("CALCUL BLOQUÉ", "Il manque une valeur, une confirmation ou un chargement conforme.", "Corriger la donnée indiquée puis confirmer de nouveau."),
        ("HORS TABLE", "La configuration dépasse un axe publié ou utilise une correction non validée.", "Ne pas extrapoler. Revenir à la documentation applicable."),
        ("Marge indicative", "Les distances ont été calculées mais la source reste provisoire.", "Lire les réserves et obtenir la validation club."),
        ("Piste insuffisante", "Une distance retenue atteint ou dépasse la distance disponible.", "Le résultat est défavorable ; ne pas l'annuler par une hypothèse plus optimiste."),
    ], [3.6, 7.1, CONTENT_WIDTH_CM - 10.7], font_size=8.1)
    heading(doc, "8.2 Comparaison aux distances disponibles", 2)
    paragraph(doc, "La distance de roulement retenue est comparée à la TORA. La distance totale jusqu'au passage des 15 m / 50 ft est comparée à la TODA. Les deux marges doivent être strictement positives ; une marge exactement nulle est refusée.")
    if "performance_result" in images:
        add_figure(doc, images["performance_result"], "Figure 6 - Exemple HORS TABLE. Aucune distance ni conclusion de piste n'est extrapolée.", width_cm=8.8, page_break_before=True)
    heading(doc, "8.3 Informations conservées", 2)
    paragraph(doc, "Même lorsqu'un calcul est bloqué ou hors table, l'écran peut conserver l'altitude-pression, la masse ou d'autres éléments certains. Ils sont affichés pour diagnostic, sans distance extrapolée et sans verdict de piste.")

    section_intro(doc, "09", "Données locales, mode hors ligne et mises à jour", "L'application peut fonctionner sans connexion pour les fonctions déjà chargées. La météo du jour nécessite une connexion et ne doit jamais être remplacée par une valeur ancienne.")
    heading(doc, "9.1 Données enregistrées sur l'appareil", 2)
    add_bullet(doc, "Les chargements saisis sont conservés localement.")
    add_bullet(doc, "Le thème jour/nuit est conservé.")
    add_bullet(doc, "Les confirmations du jour ne restent pas valides après rechargement.")
    add_bullet(doc, "L'option humidité des performances revient désactivée pour une nouvelle ouverture.")
    heading(doc, "9.2 Mode hors ligne", 2)
    paragraph(doc, "Les écrans, les données terrain et les documents déjà mis en cache restent accessibles. Sans le relais météo, le terrain peut encore être trouvé localement et les conditions doivent être saisies manuellement. Un METAR mémorisé ne doit pas être considéré comme une observation du jour.")
    heading(doc, "9.3 Cycle AIRAC", 2)
    paragraph(doc, "Le cycle et sa date d'effet sont affichés avec les données terrain. Un cycle absent, futur ou âgé de plus de 28 jours déclenche un avertissement. La VAC, l'AIP et les NOTAM en vigueur gardent la priorité.")
    heading(doc, "9.4 Mise à jour de l'application", 2)
    paragraph(doc, "Le service worker installe une nouvelle version des fichiers lorsque le réseau est disponible. Sur un iPhone, fermer puis rouvrir l'application après une mise à jour importante. Si l'affichage reste ancien, ouvrir le site dans Safari, attendre son chargement complet puis relancer l'icône de l'écran d'accueil.")

    section_intro(doc, "10", "Incidents et reprise en saisie manuelle", "Une indisponibilité de terrain ou de météo ne doit jamais conduire à fabriquer une valeur. L'application conserve les champs modifiables pour une reprise contrôlée.")
    add_data_table(doc, ["Situation", "Comportement attendu", "Reprise"], [
        ("Terrain absent", "Altitude et pistes ne sont pas préremplies.", "Saisir depuis la VAC/AIP en vigueur."),
        ("Terrain présent, pas de METAR", "Les données physiques restent utilisables.", "Saisir QNH, température, rosée et vent après contrôle."),
        ("METAR de station voisine", "La station est signalée comme voisine.", "Comparer avec les observations au terrain avant confirmation."),
        ("METAR ancien ou incohérent", "Aucune valeur météo n'est préremplie.", "Obtenir une observation valide ou saisir manuellement."),
        ("TORA/TODA absentes", "Les champs restent vides.", "Consulter la VAC, l'AIP et les NOTAM."),
        ("Pente non publiée", "0 % est proposé et étiqueté par défaut.", "Vérifier la documentation terrain et modifier si nécessaire."),
        ("Piste mouillée", "Le résultat passe hors table.", "Employer uniquement une méthode approuvée applicable."),
    ], [4.4, 6.2, CONTENT_WIDTH_CM - 10.6], font_size=7.9)
    add_callout(doc, "Signalement", "Noter l'avion, le terrain, les valeurs saisies, le message exact et la source officielle contradictoire. Ne pas corriger une donnée de référence sans identifier sa révision et sa date d'effet.", "note")

    section_intro(doc, "ANNEXE A", "Méthodes de calcul", "Cette annexe décrit les transformations appliquées par le moteur. Elle permet de reproduire un résultat et d'identifier le point où un calcul est refusé.")
    heading(doc, "A.1 Masse et centrage", 2)
    paragraph(doc, "Pour chaque poste : moment = masse x bras. Le moment total est le moment à vide publié augmenté des moments des postes. Le centre de gravité vaut : CG = moment total / masse totale.")
    paragraph(doc, "Aucun arrondi intermédiaire n'est appliqué. Les valeurs affichées sont arrondies uniquement pour la lecture. La décision d'appartenance utilise les valeurs complètes et considère la frontière de l'enveloppe comme incluse.")
    add_data_table(doc, ["Élément", "F-HDLT / F-HDLV", "F-GGHL"], [
        ("Masse à vide", "358,8 kg", "635,2 kg"),
        ("Moment à vide", "95,64 kg.m", "201,488 kg.m"),
        ("Masse maximale", "600 kg", "1 100 kg"),
        ("Densité carburant", "0,72 kg/l", "0,72 kg/l"),
        ("Repère enveloppe", "Moment / masse", "Centrage / masse"),
    ], [4.2, 6.4, CONTENT_WIDTH_CM - 10.6])
    heading(doc, "A.2 Atmosphère standard", 2)
    paragraph(doc, "La pression au terrain est estimée depuis le QNH et l'élévation dans le modèle troposphérique ISA. L'altitude-pression est ensuite l'altitude ISA possédant cette pression. La densité sèche est calculée avec l'équation des gaz parfaits, puis convertie en altitude ISA de même densité.")
    add_data_table(doc, ["Constante", "Valeur"], [
        ("Température ISA au niveau de la mer", "15 °C"),
        ("Pression standard", "1 013,25 hPa"),
        ("Densité standard", "1,225 kg/m3"),
        ("Gradient troposphérique", "0,0065 K/m"),
        ("Gravité", "9,80665 m/s2"),
        ("Constante de l'air sec", "287,05 J/(kg.K)"),
        ("Constante de la vapeur d'eau", "461,495 J/(kg.K)"),
    ], [9.0, CONTENT_WIDTH_CM - 9.0])
    heading(doc, "A.3 Air humide", 2)
    paragraph(doc, "Le point de rosée est converti en pression de vapeur par une relation de Magnus. La pression totale est séparée en pression partielle d'air sec et pression partielle de vapeur d'eau. Les deux densités partielles sont additionnées, puis la densité obtenue est convertie en altitude-densité ISA.")
    paragraph(doc, "Ce calcul est un calcul atmosphérique général. Dans le module Performances seulement, une altitude-pression sèche équivalente est calculée puis soumise à la table. La distance la plus restrictive entre le calcul standard et cette comparaison est retenue.")
    heading(doc, "A.4 Interpolation SportStar RTC", 2)
    paragraph(doc, "Le moteur calcule l'altitude-pression et l'écart à la température ISA. Il effectue une interpolation bilinéaire entre les deux altitudes et les deux écarts ISA encadrant le point. Les distances « dur » et « herbe », roulement et 15 m / 50 ft, restent quatre colonnes distinctes.")
    add_data_table(doc, ["Axe", "Domaine actif"], [
        ("Altitude-pression", "0 à 10 000 ft"),
        ("Écart ISA", "ISA-10 à ISA+20 °C"),
        ("Masse", "Valeurs publiées conservées jusqu'à 600 kg ; aucun gain sous la masse de référence"),
    ], [6.0, CONTENT_WIDTH_CM - 6.0])
    heading(doc, "A.5 Interpolation DR400/180", 2)
    paragraph(doc, "Le moteur effectue une interpolation trilinéaire entre masse, altitude-pression et écart ISA. Les points publiés sont 900 et 1 100 kg, 0, 4 000 et 8 000 ft, puis ISA-20, ISA et ISA+20 °C.")
    paragraph(doc, "Sous 900 kg, la ligne 900 kg est retenue sans extrapolation vers le bas. Cette convention est conservatrice par rapport à la table utilisée. Au-dessus de 1 100 kg, le résultat est hors table.")
    heading(doc, "A.6 Corrections et ordre d'application", 2)
    add_step(doc, "Table", "interpoler séparément le roulement et la distance totale.")
    add_step(doc, "Vent", "appliquer, pour le DR400 seulement, le facteur de vent de face publié entre 0 et 30 kt. Aucun vent arrière n'est corrigé.")
    add_step(doc, "Humidité facultative", "répéter le calcul à l'altitude-pression sèche équivalente et conserver la valeur la plus restrictive.")
    add_step(doc, "Marge additionnelle", "multiplier la distance retenue par (1 + marge/100).")
    add_step(doc, "Arrondi", "arrondir chaque distance au mètre supérieur.")
    add_step(doc, "Comparaison", "comparer le roulement à la TORA et la distance totale à la TODA.")
    heading(doc, "A.7 Facteurs de vent DR400", 2)
    add_data_table(doc, ["Vent de face", "Facteur"], [("0 kt", "1,00"), ("10 kt", "0,81"), ("20 kt", "0,67"), ("30 kt", "0,56")], [8.0, CONTENT_WIDTH_CM - 8.0])
    paragraph(doc, "Les facteurs intermédiaires sont interpolés linéairement. Au-delà de 30 kt, le calcul est hors table. Aucun facteur n'est transféré au SportStar RTC faute d'identification certaine de la série de corrections.")
    heading(doc, "A.8 Refus d'extrapolation", 2)
    paragraph(doc, "Si l'altitude-pression, l'écart ISA, la masse ou une correction sort du domaine publié, aucune distance n'est calculée à partir d'une prolongation mathématique de la table. Le message HORS TABLE conserve les informations certaines mais retire le verdict de piste.")

    section_intro(doc, "ANNEXE B", "Données intégrées par avion", "Les tableaux suivants permettent de contrôler les constantes utilisées par la version 0.9.")
    heading(doc, "B.1 Evektor SportStar - F-HDLT et F-HDLV", 2)
    add_data_table(doc, ["Poste", "Bras", "Valeur / limite intégrée"], [
        ("Pilote", "0,545 m", "Masse positive ; limite individuelle à vérifier"),
        ("Copilote", "0,545 m", "Équipage total : domaine du tableau jusqu'à 220 kg"),
        ("Bagages", "1,083 m", "25 kg maximum selon la fiche transmise"),
        ("Carburant", "0,680 m", "120 l maximum ; plein prérempli"),
    ], [6.0, 3.4, CONTENT_WIDTH_CM - 9.4])
    paragraph(doc, "Sources : fiche de pesée F-HDLT du 28/01/2025 et fiche club « Calcul masse et centrage EVSS F-HDLT ». Les données sont appliquées provisoirement à F-HDLV.")
    paragraph(doc, "Sommets de l'enveloppe moment / masse : (95 ; 375), (120 ; 375), (150 ; 400), (180 ; 450), (240 ; 600), (225 ; 600). Cette numérisation doit être vérifiée sur le manuel approuvé.")
    heading(doc, "B.2 Robin DR400/180 - F-GGHL", 2)
    add_data_table(doc, ["Poste", "Bras", "Valeur / limite intégrée"], [
        ("Pilote et passager avant", "0,410 m", "Limites individuelles non fournies dans le scan"),
        ("Passagers arrière", "1,190 m", "Limite du poste à vérifier"),
        ("Bagages", "1,900 m", "Limite de soute à vérifier"),
        ("Carburant principal", "1,120 m", "110 l total ; 10 l inutilisables déjà en masse à vide"),
        ("Carburant ailes", "0,100 m", "2 x 40 l ; plein prérempli à 80 l"),
    ], [6.2, 3.3, CONTENT_WIDTH_CM - 9.5])
    paragraph(doc, "Source : rapport de pesée F-GGHL du 09/06/2015. L'enveloppe intégrée correspond à la catégorie Normale. La catégorie Utilitaire visible sur la fiche n'est pas évaluée.")
    paragraph(doc, "Sommets centrage / masse : (0,205 ; 635,2), (0,205 ; 750), (0,428 ; 1 100), (0,564 ; 1 100), (0,564 ; 635,2).")
    heading(doc, "B.3 Tables de performances", 2)
    add_data_table(doc, ["Avion", "Source active", "Statut"], [
        ("F-HDLT / F-HDLV", "Clichés « EVEKTOR SPORTSTAR RTC - Page 60 »", "Série et conditions générales à valider"),
        ("F-GGHL", "Manuel générique DR400/180, édition 13, table page 5.2", "Non spécifique à l'avion ; applicabilité à confirmer"),
    ], [4.2, 8.4, CONTENT_WIDTH_CM - 12.6], font_size=8.1)
    heading(doc, "B.4 Données terrain", 2)
    paragraph(doc, "Le fichier local est généré à partir du catalogue SIA AD 1.3 / AIXM du cycle AIRAC 08/26. Il contient 432 aérodromes métropolitains et 1 349 entrées directionnelles. LFRQ et LFMW disposent en plus de distances déclarées intégrées depuis leur documentation courante.")
    paragraph(doc, "À titre de contrôle : LFRQ, altitude 297 ft, QFU 09/27, asphalte, 2 150 x 45 m ; TORA/TODA 09 : 2 150/2 150 m et 27 : 2 113/2 113 m. LFMW, altitude 553 ft, QFU 11/29, piste revêtue 810 x 30 m ; TORA/TODA 810/810 m.")

    section_intro(doc, "ANNEXE C", "Réserves d'emploi et points à valider", "Cette liste accompagne la version de travail. Elle doit être soldée ou maintenue explicitement avant une diffusion opérationnelle.")
    heading(doc, "C.1 Réserves propres aux avions", 2)
    add_bullet(doc, "Confirmer que la pesée, l'enveloppe et les performances de F-HDLT peuvent être appliquées à F-HDLV.")
    add_bullet(doc, "Confirmer l'identification de la série SportStar RTC et ses conditions générales : masse, volets, puissance, technique, hélice et état de piste.")
    add_bullet(doc, "Confirmer le manuel, la révision, l'hélice, les carénages et les suppléments applicables à F-GGHL.")
    add_bullet(doc, "Vérifier les limites individuelles de sièges, de banquette arrière et de soute absentes des scans actuels.")
    add_bullet(doc, "Contrôler les sommets des enveloppes numérisées et décider du traitement de la catégorie Utilitaire du DR400.")
    heading(doc, "C.2 Réserves relatives aux performances", 2)
    add_bullet(doc, "Aucune extrapolation ne doit être autorisée hors des axes publiés.")
    add_bullet(doc, "Les pistes mouillées restent hors table tant qu'un facteur applicable n'a pas été validé.")
    add_bullet(doc, "Le SportStar RTC n'applique actuellement aucune correction de vent ou de pente issue des séries photographiées non identifiées.")
    add_bullet(doc, "Le DR400 refuse le vent arrière, la pente montante et un vent de face supérieur à 30 kt faute de correction couverte.")
    add_bullet(doc, "La comparaison avec l'air humide est une estimation complémentaire ; elle n'est pas une méthode publiée dans les manuels utilisés.")
    add_bullet(doc, "La marge additionnelle est réglée à 0 % par défaut. Une valeur club ne doit être intégrée qu'après décision formelle.")
    heading(doc, "C.3 Réserves relatives aux terrains et à la météo", 2)
    add_bullet(doc, "Définir et appliquer une procédure de mise à jour du jeu SIA tous les 28 jours.")
    add_bullet(doc, "Ne pas confondre longueur physique, TORA, TODA, ASDA et LDA.")
    add_bullet(doc, "Traiter 0 % de pente comme une valeur par défaut lorsque la pente n'est pas publiée.")
    add_bullet(doc, "Une station METAR voisine ne décrit pas automatiquement le terrain demandé.")
    add_bullet(doc, "La proposition de QFU ne remplace pas la piste en service ni les contraintes opérationnelles.")
    add_bullet(doc, "Les terrains de montagne, altiports et altisurfaces nécessitent leur documentation et leur réglementation propres.")
    heading(doc, "C.4 Critère de publication", 2)
    paragraph(doc, "Après validation, chaque donnée doit pouvoir être rattachée à un avion, une source, une révision et une date d'effet. Une nouvelle pesée ou une nouvelle révision de manuel doit invalider explicitement l'ancienne configuration. Une modification silencieuse n'est pas acceptable.")
    add_callout(doc, "Statut actuel", "La version 0.9 du présent manuel décrit fidèlement l'application, mais ne l'approuve pas. La validation doit être réalisée par les personnes désignées par l'aéroclub.", "attention")

    section_intro(doc, "ANNEXE D", "Glossaire", "Les termes employés dans l'interface sont repris ci-dessous dans leur sens opérationnel pour l'application.")
    add_data_table(doc, ["Terme", "Définition"], [
        ("AIP", "Publication d'information aéronautique officielle."),
        ("AIRAC", "Cycle de mise à jour planifiée des données aéronautiques."),
        ("Altitude-densité", "Altitude ISA correspondant à la densité de l'air calculée."),
        ("Altitude-pression", "Altitude ISA correspondant à la pression statique calculée au terrain."),
        ("ASDA", "Distance utilisable pour l'accélération-arrêt."),
        ("CG", "Centre de gravité."),
        ("Écart ISA", "Température extérieure moins température ISA à l'altitude-pression."),
        ("LDA", "Distance utilisable à l'atterrissage."),
        ("METAR", "Observation météorologique d'aérodrome."),
        ("QFU", "Orientation de piste identifiée ; le calcul utilise le relèvement vrai publié lorsqu'il existe."),
        ("QNH", "Pression ramenée au niveau moyen de la mer."),
        ("TAF", "Prévision d'aérodrome, affichée mais non utilisée dans le calcul instantané."),
        ("TODA", "Distance utilisable au décollage, prolongement dégagé compris lorsqu'il est publié."),
        ("TORA", "Distance de roulement utilisable au décollage."),
        ("VAC", "Carte d'approche et d'atterrissage à vue."),
    ], [4.6, CONTENT_WIDTH_CM - 4.6], font_size=8.2)
    heading(doc, "Sources de référence", 2)
    p = paragraph(doc, "SIA - produits numériques et Atlas VAC : ", size=8.8)
    add_hyperlink(p, "sia.aviation-civile.gouv.fr", "https://www.sia.aviation-civile.gouv.fr/")
    p = paragraph(doc, "Aviation Weather Center - Data API : ", size=8.8)
    add_hyperlink(p, "aviationweather.gov/data/api", "https://aviationweather.gov/data/api/")
    p = paragraph(doc, "Copie publique du manuel générique DR400/180 utilisée : ", size=8.8)
    add_hyperlink(p, "manualzilla.com/doc/6382928", "https://manualzilla.com/doc/6382928/manuel-de-vol-robin-dr400-180")
    p = paragraph(doc, "Fiche de type EASA DR 200/300/400 : ", size=8.8)
    add_hyperlink(p, "easa.europa.eu - EASA.A.367", "https://www.easa.europa.eu/en/document-library/type-certificates/aircraft-cs-25-cs-22-cs-23-cs-vla-cs-lsa/easaa367-ceapr-dr-200")
    paragraph(doc, "Les scans propres au club restent consultables depuis la page « Méthodes & sources » de l'application.", size=8.8, color=MUTED)

    # Ensure a clean last paragraph without an accidental blank page.
    doc.core_properties.title = "Outils de vol - Manuel utilisateur et méthodes de calcul"
    doc.core_properties.subject = "Manuel utilisateur - Aéroclub Jean-Doudiès"
    doc.core_properties.author = "Aéroclub Jean-Doudiès"
    doc.core_properties.comments = "Version de travail 0.9 - validation club requise"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--toc-json", type=Path)
    args = parser.parse_args()
    toc_pages = {}
    if args.toc_json and args.toc_json.exists():
        toc_pages = json.loads(args.toc_json.read_text(encoding="utf-8"))
    path = create_manual(toc_pages)
    print(path)


if __name__ == "__main__":
    main()
